from docx import Document
import markdown2
from bs4 import BeautifulSoup

def ai_to_structured_docx(ai_text, output_file="output.docx"):
    html = markdown2.markdown(ai_text, extras=["tables"])
    soup = BeautifulSoup(html, 'html.parser')
    doc = Document()

    for element in soup.children:
        if element.name == 'h1':
            # Uses Word's "Heading 1" style
            doc.add_heading(element.text, level=1) 
        elif element.name == 'h2':
            doc.add_heading(element.text, level=2) # "Heading 2"
        elif element.name == 'p':
            p = doc.add_paragraph()
            # Check for bold/italic inside the paragraph
            for content in element.contents:
                if content.name == 'strong':
                    run = p.add_run(content.text)
                    run.bold = True   # Word's "Strong" style
                elif content.name == 'em':
                    run = p.add_run(content.text)
                    run.italic = True # Word's "Emphasis" style
                else:
                    p.add_run(content.text) # "Normal" style
        elif element.name == 'ul':
            for li in element.find_all('li'):
                doc.add_paragraph(li.text, style='List Bullet')
        elif element.name == 'table':
            # Add table with "Table Grid" style
            rows = element.find_all('tr')
            if rows:
                table = doc.add_table(rows=len(rows), cols=len(rows[0].find_all(['th', 'td'])))
                table.style = 'Table Grid'
                # ... fill cells ...

    doc.save(output_file)
    print(f"✅ Document saved as {output_file}")
    
ai_response = """
# Sales Report

Here is the quarterly data:

| Quarter | Revenue | Profit |
|---------|---------|--------|
| Q1      | $100K   | $20K   |
| Q2      | $150K   | $35K   |
| Q3      | $200K   | $50K   |
"""

ai_to_structured_docx(ai_response, "report.docx")